from flask import Flask, render_template, request, jsonify
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import AzureOpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document as LangChainDocument
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from doc2docx import convert
from werkzeug.utils import secure_filename
from uuid import uuid4
from dotenv import load_dotenv
import os
import time
import json

load_dotenv()
os.environ["AZURE_OPENAI_ENDPOINT"] = os.getenv("AZURE_OPENAI_ENDPOINT")
os.environ["AZURE_OPENAI_API_KEY"] = os.getenv("AZURE_OPENAI_API_KEY")

vector_store = Chroma(
    collection_name = "fareastonHR",
    embedding_function = AzureOpenAIEmbeddings(
        model="text-embedding-ada-002", 
        openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION")
    ),
    persist_directory="../db/fareastonHRupdate2"
)

app = Flask(__name__)

documents=[]

def get_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_from_docx(docx_path):
    doc = DocxDocument(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

data_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data')
if not os.path.exists(data_folder):
    os.makedirs(data_folder)


@app.route('/')
def index():
    return render_template('manager.html')

extension_list=[]
require_extension=['.docx','.pdf','.doc']
@app.route('/upload_file', methods=['POST'])
def upload_file():
    folder_name = request.form.get('folderName')
    filename_list = []
    file_chunk_map = {}
    json_path = os.path.join(os.path.dirname(__file__), 'file_chunk_ids.json')
    if os.path.exists(json_path):
        with open(json_path, 'r', encoding='utf-8') as f:
            file_chunk_map = json.load(f)
             
    for root, dirs, files in os.walk(folder_name):
        for file in files:
            file_path = os.path.relpath(os.path.join(root, file), start=data_folder)
            filename_list.append(file_path)

    if 'filename' in request.files:
        files = request.files.getlist('filename')
        documents = []
        text_splitter = CharacterTextSplitter(
            separator="\n",
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        for file in files:
            if file.filename != '':
                filename = file.filename
                file_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', folder_name, filename)
                _, extension = os.path.splitext(file_path)
                if filename not in filename_list and (extension in require_extension):
                    file.save(file_path)
                    extension_list.append(extension)

                    if filename.endswith('.pdf'):
                        text = get_text_from_pdf(file_path)
                    elif filename.endswith('.docx'):
                        text = get_text_from_docx(file_path)
                    elif filename.endswith('.doc'):
                        output_path = os.path.splitext(file_path)[0] + "_output.docx"
                        if not os.path.exists(output_path):
                            convert(file_path, output_path)
                            os.remove(file_path)
                        text = get_text_from_docx(output_path)
                    chunks = text_splitter.split_text(text)
                    uuids = []
                    for chunk in chunks:
                        uuid = str(uuid4())
                        doc = LangChainDocument(page_content=chunk, metadata={"source": extension}, id=uuid)
                        documents.append(doc)
                        uuids.append(uuid)

                    file_chunk_map[filename] = uuids 
        if documents:
            vector_store.add_documents(documents=documents, ids=[doc.id for doc in documents])
            with open('file_chunk.json', 'w', encoding='utf-8') as f:
                json.dump(file_chunk_map, f, ensure_ascii=False)
            return jsonify({'filenames': filename_list, 'extensions': extension_list})
        
    return 'No file uploaded.'

@app.route('/deleteText', methods=['POST'])
def deleteFile():
    data = request.json
    filename = data['filename']
    folderName = data['folderName']
    file_path=os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', folderName, filename)
    file = file_path
    with open('file_chunk_ids.json', 'r', encoding='utf-8') as f:
        file_chunk_map = json.load(f)

    if filename in file_chunk_map:
        uuids = file_chunk_map[filename]
        for uuid in uuids:
            vector_store.delete(ids = uuid)
        del file_chunk_map[filename]
        
        with open('file_chunk_ids.json', 'w', encoding='utf-8') as f:
            json.dump(file_chunk_map, f, ensure_ascii=False)

    os.remove(file)
    return jsonify({'message': 'File deleted'})

@app.route('/getFiles', methods=['POST'])
def get_all_files():
    request_data = request.get_json()
    folder_name = request_data.get('folderName', '')
    filelist = []

    data_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', folder_name)

    for root, dirs, files in os.walk(data_folder):
        for file in files:
            file_path = os.path.relpath(os.path.join(root, file), start=data_folder)
            filelist.append(file_path)
    
    return jsonify(files=filelist)

@app.route('/getFolders')
def get_folders():
    folder_path = '../data' 
    folders = []
    for entry in os.listdir(folder_path):
        if os.path.isdir(os.path.join(folder_path, entry)):
            folder_info = {
                'name': entry,
                'time': time.ctime(os.path.getmtime(os.path.join(folder_path, entry)))
            }
            folders.append(folder_info)
    return jsonify(folders=folders)

if __name__ == '__main__':
    app.run(debug=True, port = 4999)

